import re
import os

with open('generate_reports.py', 'r', encoding='utf-8') as f:
    text = f.read()

# 1. Provide doc=None arguments and add is_combined flag
text = text.replace('def create_proje2_report(doc=None):', 'def create_proje2_report(doc=None, is_combined=False):')
text = text.replace('def create_proje5_report(doc=None):', 'def create_proje5_report(doc=None, is_combined=False):')

# 2. Skip cover/toc based on is_combined flag
text = text.replace(
'''    add_cover_page(doc, "Veritabanı Yedekleme ve Felaketten Kurtarma Planı", "2")
    add_toc(doc)''',
'''    if not is_combined:
        add_cover_page(doc, "Veritabanı Yedekleme ve Felaketten Kurtarma Planı", "2")
        add_toc(doc)
    else:
        doc.add_page_break()
        doc.add_heading("PROJE 2: VERİTABANI YEDEKLEME VE FELAKETTEN KURTARMA PLANI", level=0)'''
)

text = text.replace(
'''    add_cover_page(doc, "Veri Temizleme ve ETL Süreçleri Tasarımı", "5")
    add_toc(doc)''',
'''    if not is_combined:
        add_cover_page(doc, "Veri Temizleme ve ETL Süreçleri Tasarımı", "5")
        add_toc(doc)
    else:
        doc.add_page_break()
        doc.add_heading("PROJE 5: VERİ TEMİZLEME VE ETL SÜREÇLERİ TASARIMI", level=0)'''
)

# 3. Replace the __main__ execution block entirely
new_main = '''if __name__ == '__main__':
    from docx import Document
    
    # KOMBİNE RAPOR
    master_doc = Document()
    set_doc_style(master_doc)
    
    # Margin
    section = master_doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # A single cover page and A single TOC
    add_cover_page(master_doc, "Veritabanı Projeleri İncelemesi (Proje 2 & Proje 5)", "2 ve 5")
    add_toc(master_doc)
    
    # Generate contents inside the master doc without their own cover pages
    create_proje2_report(doc=master_doc, is_combined=True)
    create_proje5_report(doc=master_doc, is_combined=True)
    
    # Save the strictly combined file
    combined_path = os.path.join(os.path.dirname(__file__), 'BLM4522_Tekli_Final_Rapor.docx')
    master_doc.save(combined_path)
    
    print("\\n=============================================")
    print(f"  --> MÜKEMMEL TEKLİ RAPOR OLUŞTURULDU: {combined_path}")
    print("=============================================")
'''

# We use regex to substitute everything from if __name__ == '__main__': to EOF
text = re.sub(r"if __name__ == '__main__':.*", new_main, text, flags=re.DOTALL)

with open('generate_reports.py', 'w', encoding='utf-8') as f:
    f.write(text)
