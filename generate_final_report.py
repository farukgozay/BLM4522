from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, ns
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

STUDENT = {
    "ad": "Muhammed Faruk GÖZAY",
    "no": "22290673",
    "bolum": "Bilgisayar Mühendisliği",
    "ders": "BLM4522 Ağ Tabanlı Paralel Dağıtım Sistemleri",
    "hoca": "Öğr. Gör. Enver BAĞCI",
    "universite": "Ankara Üniversitesi",
    "fakulte": "Mühendislik Fakültesi",
    "github": "https://github.com/farukgozay/BLM4522.git",
    "tarih": "Nisan 2026"
}

gorsel_counter = 1

def add_image_placeholder(doc, caption, instruction):
    global gorsel_counter
    # Bolca boşluk bırakıyoruz
    doc.add_paragraph("\n")
    doc.add_paragraph("\n")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"--- [ ✂ BURADAKİ KIZIL YAZILARI SİLİP EKRAN GÖRÜNTÜSÜ YAPIŞTIRIN ] ---\n\nSS YÖNERGESİ: {instruction}\n\n---------------------------------")
    r.font.color.rgb = RGBColor(255, 0, 0)
    r.bold = True
    
    doc.add_paragraph("\n")
    doc.add_paragraph("\n")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Görsel {gorsel_counter}: {caption}")
    r2.bold = True
    r2.font.size = Pt(10)
    
    doc.add_paragraph("\n")
    
    gorsel_counter += 1

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(ns.qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')
    instrText.text = " PAGE "

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(ns.qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(ns.qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color, qn('w:val'): 'clear'
    })
    shading.append(shd)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2E4057")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0F4F8")
    doc.add_paragraph("\n")
    return table

def add_cover_page(doc, proje_adi):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(STUDENT["universite"])
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(STUDENT["fakulte"])
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(STUDENT["bolum"])
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(STUDENT["ders"])
    r.bold = True
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(proje_adi)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    for _ in range(4):
        doc.add_paragraph()

    info_lines = [
        ("Öğrenci", f"{STUDENT['ad']}"),
        ("Öğrenci No", STUDENT["no"]),
        ("Dersin Hocası", STUDENT["hoca"]),
        ("GitHub", STUDENT["github"]),
        ("Tarih", STUDENT["tarih"]),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}: ")
        r.bold = True
        r.font.size = Pt(11)
        r = p.add_run(value)
        r.font.size = Pt(11)

    doc.add_page_break()

def add_toc(doc):
    doc.add_heading('İçindekiler', level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fldChar1 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._element.append(fldChar1)
    run2 = p.add_run()
    instrText = run2._element.makeelement(qn('w:instrText'), {})
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instrText)
    run3 = p.add_run()
    fldChar2 = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    run3._element.append(fldChar2)
    run4 = p.add_run("(ÖNEMLİ BİLGİ: Python sayfa katsayısını bilemeyeceği için sayfa numaraları şu an rastgeledir. Doğru sayfa numaralarını görmek için WORD'DE BU YAZIYA SAĞ TIKLAYIP 'Alanı Güncelleştir' -> 'Tüm tabloyu güncelleştir' seçeneğini TIKLAYINIZ!)")
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run4.font.size = Pt(9)
    run5 = p.add_run()
    fldChar3 = run5._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run5._element.append(fldChar3)
    doc.add_page_break()

def set_doc_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
        hs.font.name = 'Calibri'

def add_code_block(doc, code, title=None):
    if title:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10)
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    pf = p.paragraph_format
    pf.left_indent = Cm(0.5)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    r = p.add_run(code)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    set_cell_shading_paragraph(p, "F5F5F5")
    doc.add_paragraph("\n")

def set_cell_shading_paragraph(paragraph, color):
    pPr = paragraph._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:fill'): color, qn('w:val'): 'clear'
    })
    pPr.append(shd)

def generate_unified_report():
    doc = Document()
    set_doc_style(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    add_cover_page(doc, "PROJE 2 & PROJE 5 KOMBİNE RAPORU")
    add_toc(doc)

    # ----------------------------------------------------
    # PROJE 2 BÖLÜMÜ
    # ----------------------------------------------------
    doc.add_heading('PROJE 2: Veritabanı Yedekleme ve Felaketten Kurtarma Planı', level=1)

    doc.add_heading('1. Giriş', level=2)
    doc.add_paragraph('Bu projede, bir e-ticaret veritabanı üzerinde yedekleme stratejileri tasarlanmış ve felaketten kurtarma planları uygulanmıştır. PostgreSQL veritabanı yönetim sistemi kullanılarak tam yedekleme, WAL tabanlı sürekli arşivleme, zamanlayıcı ile otomatik yedekleme ve çeşitli felaket senaryolarında veri kurtarma işlemleri gerçekleştirilmiştir.')
    doc.add_paragraph('Veri yönetimi, günümüz e-ticaret platformlarının kesintisiz çalışması (high availability) için vazgeçilmez bir unsurdur. Özellikle SQL Server ve PostgreSQL gibi devasa veri yığınlarına hizmet veren sistemlerde sadece donanım güvenliği değil, veritabanının kendi iç felaket senaryosu planlaması da kusursuz hesaplanmalıdır. Bu çalışma, sektörel standartlarda tasarlanmış 7 farklı tablo ile gerçek zamanlı sipariş ve müşteri yönetimini simüle eden bir ortam üzerinde oluşturulmuştur.')
    
    doc.add_heading('1.1. Projenin Amacı ve Mimari Metodoloji', level=3)
    doc.add_paragraph('Veritabanı yönetiminde en kritik konulardan biri veri kaybına karşı önlem almaktır. Bu proje kapsamında aşağıdaki temel hedefler gerçekleştirilmiştir:')
    bullets = [
        'Tam, artık ve fark yedekleme stratejilerinin sistematik bir mimari tasarımla planlanması ve uygulanması.', 
        'Zamanlayıcı tabanlı, insan müdahalesine gerek duymayan (No-Touch) otomatik yedekleme sisteminin kurulması.', 
        'Felaket senaryolarının (Drop Table, Delete from Without Where, Hardware Failure) simüle edilmesi ve rollback testleri.', 
        'Yedeklerin kriptolojik bütünselliği için (MD5 Checksum) doğrulama mekanizmalarının entegre edilmesi.', 
        'Veritabanında her yapılan operasyonu ve hatayı kayıt altına alan (Audit Log) izleme şeması.'
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
        
    doc.add_paragraph('Bu projede hedef sadece teknik olarak bir veriyi zip formatına sıkıştırmak değil; veri tabanını belirli bir saniyedeki anlık zaman damgasına (Point in Time Recovery) kadar hiç firesiz geri yükleyebilmektir.')

    doc.add_heading('2. Kullanılan Teknolojiler', level=2)
    add_styled_table(doc, ['Teknoloji', 'Sürüm', 'Kullanım Alanı'], [['PostgreSQL', '16.x/18.x', 'Veritabanı yönetim sistemi'], ['pgAdmin 4', 'Güncel', 'Grafik arayüz yönetimi ve sorgu aracı'], ['pg_dump / pg_restore', 'Dahili', 'Yedekleme ve geri yükleme'], ['pg_basebackup', 'Dahili', 'Base backup (WAL için sürekli akış)'], ['Windows Task Scheduler', 'Dahili', 'Otomatik zamanlama aracı'], ['Batch Script (.bat)', '-', 'Otomasyon scriptleri yazımı']])

    doc.add_page_break()

    doc.add_heading('3. Veritabanı Tasarımı', level=2)
    doc.add_paragraph('Projede "eticaret_db" adında bir e-ticaret veritabanı tasarlanmıştır. Bu veritabanı müşteriler, ürünler, siparişler, ödemeler ve yedekleme loglarını içeren birbiriyle bütünleşik 7 tablodan oluşmaktadır.')

    doc.add_heading('3.1. E-Ticaret Veritabanı Tablo Yapıları ve Açıklamaları', level=3)
    add_styled_table(doc, ['Tablo Adı', 'Açıklama', 'Kayıt Sayısı'], [['musteriler', 'Müşteri bilgileri: ad, soyad ve iletişim bilgisi', '50'], ['kategoriler', 'Ürün kategorileri', '10'], ['urunler', 'Ürün bilgileri ve temel stok durumları', '30'], ['siparisler', 'Sipariş ana gövdesi (Tutar ve referans)', '100'], ['siparis_detaylari', 'Siparişte bulunan kalemlerin ürün alt tipleri', '~200'], ['odemeler', 'Banka ödeme durumu ve bekleyen işlemler', '~60'], ['yedekleme_log', 'Tarihsel yedekleme işlemlerinin boyut logları', 'Dinamik']])

    add_image_placeholder(doc, "eticaret_db Tablo Görünümü ve Şema", "pgAdmin 4'ü açın. eticaret_db > Schemas > public > Tables dizinini tam açık gösterecek şekilde soldaki menünün tam ekran görüntüsünü alıp buraya yapıştırın.")

    doc.add_heading('4. Yedekleme Stratejileri', level=2)
    doc.add_paragraph('PostgreSQL\'in pg_dump aracı kullanılarak veritabanının tam yedeği alınmıştır. Bu adımda formatı özellikle ".dump" uzantısı kullanılarak Customs ayarında tasarlanmıştır ki sonradan sadece belirli bir tabloyu bile geri yükleme esnekliğine sahip olalım.')
    add_code_block(doc, 'pg_dump -Fc -v -f "C:\\pg_backups\\full\\full_backup.dump" eticaret_db\npg_dump --format=plain -f "C:\\pg_backups\\full\\full_backup.sql" eticaret_db', 'Tam Yedekleme Komutu (Full Backup):')
    doc.add_paragraph('Windows Task Scheduler kullanılarak yedekleme işlemleri otomatize edilmiştir. Eski yedek dosyalarını temizlemek için 7 günden eski dosyaların otomatik silinmesi sağlanmıştır.')

    doc.add_heading('5. Felaketten Kurtarma Senaryoları', level=2)
    doc.add_paragraph('Veri sızıntısından veya büyük kullanıcı (Human Error) hatalarından anlık şekilde dönmek çoğu zaman tam veri yüklemekten daha hızlıdır. Transaction ve SAVEPOINT mekanizması kullanılarak kazara silinen verilerin anında geri alınması sağlanmıştır. Aşağıda BEGIN, SAVEPOINT ve ROLLBACK TO işleminin uygulamalı kodu mevcuttur:')
    add_code_block(doc, 'BEGIN;\nSAVEPOINT once_savepoint;\nDELETE FROM musteriler WHERE sehir = \'İstanbul\';\nROLLBACK TO SAVEPOINT once_savepoint;\nCOMMIT;')

    add_image_placeholder(doc, "pgAdmin Query Tool Çalışan Rollback Hata Kurtarma Kodu", "pgAdmin Query Tool'da üstteki BEGIN; DELETE FROM vb.. komutu yazın. Seçip bir kez F5'e basarak 'Query returned successfully' yazan ekranın alt kısımla birlikte SS'ini alın.")

    doc.add_page_break()

    # ----------------------------------------------------
    # PROJE 5 BÖLÜMÜ
    # ----------------------------------------------------
    doc.add_heading('PROJE 5: Veri Temizleme ve ETL Süreçleri Tasarımı', level=1)

    doc.add_heading('1. Bütünleşik ETL Mimarisine Giriş', level=2)
    doc.add_paragraph('Bu projede, bir e-ticaret sistemine ait ham verilerin (Dirty Data) ETL (Extract, Transform, Load) süreçleri ile tamamen hijyenik ve kurumsal bir standarda çekilmesi, dönüştürülmesi ve kısıtlamalı hedef tablolara yüklenmesi gerçekleştirilmiştir. Gerçek hayatta kullanıcılar, eski veri sistemleri veya API\'ler, veri tabanınıza eksik veya biçimi bozuk yüzlerce hatalı format sokabilir. Eğer ETL boru hattı (Pipeline) olmazsa veritabanınız çöplüğe döner.')

    doc.add_heading('1.1. Veri Kalitesinin Temel Prensipleri ve Önemi', level=3)
    doc.add_paragraph('Gerçek dünyada veri kaynakları çoğu zaman tutarsızdır ve aşağıdaki hususlar bu süreç için çok değerlidir:')
    bullets = [
        'Dışarıdan gelen devasa CSV/JSON/API verilerini "Extract" adımı ile önce tampon bir bölgeye almak (Staging).', 
        'O bölgedeki her bir kelime, regex (düzenli ifade) komutlarıyla ve string format testleriyle "Transform" (Dönüşüm) adımından geçirmek.', 
        'Temizlenen verilerin asla Foreign Key (Yabancı Anahtar) referans boşluğuna düşmemesi için Load adımını tasarlamak.', 
        'Hata sızdırmazlığını ölçmek için Veri Kalitesi Raporları çıkararak yönetime dashboard sunmak.'
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_page_break()

    doc.add_heading('2. ETL Mimarisi Şemaları', level=2)
    doc.add_paragraph('Projede iş mantığı gereği veri üç farklı evreden geçmektedir: Ham veri, Log ve Temiz veri. Mimari şemalar şu şekildedir:')
    add_styled_table(doc, ['Şema (Schema) Adı', 'Amaç (Açıklama)', 'Kısıtlama ve Özellik'], [['staging', 'Ham verilerin dışarıdan okunduğu alan', 'Hiç bir kısıtlama yok, tüm veriler TEXT'], ['hedef', 'Temiz ve son onayı almış verilerin alanı', 'Zorunlu PK, FK, CHECK kısıtlamaları var'], ['etl_log', 'Süreç istatistikleri ve tespit edilen hatalar', 'Dinamik hata kaydı tutar, analitik için kullanılır']])

    add_image_placeholder(doc, "etl_db Veritabanı ve Şema Hiyerarşisi", "pgAdmin 4'ü açın. etl_db veritabanını genişletin. Schemas altındaki 'etl_log', 'hedef' ve 'staging' yazan 3 ayrı katmanı okunaklı gösterecek şekilde SS alın.")

    doc.add_heading('3. Karşılaşılan Ham Veri Sorunları (Dirty Data)', level=2)
    doc.add_paragraph('Ham müşteri, sipariş ve ürün verilerini incelerken sistemin karşı karşıya kaldığı temel problemleri sınıflandırmak gerekmektedir:')
    add_styled_table(doc, ['Sorun Tipi', 'Hangi Tablolarda Bulundu?', 'Tarafımca Uygulanan SQL Çözümü'], [['Duplike aynı email kayıtları', 'Müşteri, Sipariş, Ürün', 'ROW_NUMBER() OVER(PARTITION BY) fonksiyonu ile temizlendi'], ['Matematiksel Negatif/Sıfır Fiyat', 'Ürün', 'CASE WHEN ifadeleri ve ABS() kullanıldı'], ['Büyük/Küçük Harf Yazım Hataları', 'Ürün, Müşteri', 'LOWER(), INITCAP() kullanıldı'], ['Tarih formatları tutarsızlığı (Eğik Çizgi)', 'Siparişler, Müşteriler', 'REGEX ile format belirlenip TO_DATE(text, format) uygulandı']])

    doc.add_page_break()

    doc.add_heading('4. ETL Aşamalarının Çalışma Prensibi', level=2)
    
    doc.add_heading('4.1. Extract (Çıkarma) Prensibi', level=3)
    doc.add_paragraph('PostgreSQL\'in COPY veya \\copy komutu kullanılarak Excel (CSV) dosyalarından staging şemasındaki tablolara ham veri indirilmiştir. Staging tablolarındaki tüm sütunlar bilerek "TEXT" veri tipindedir ki veritabanı kilitlendi veya Cast Error (dönüştürme hatası) fırlatmasın. Çıkarma aşaması %100 sorunsuz bir depolama anıdır.')
    add_code_block(doc, "\\copy staging.raw_musteriler FROM 'raw_customers.csv' WITH (FORMAT csv, HEADER true, ENCODING 'UTF8')", 'Basit Bir COPY Uygulaması:')

    add_image_placeholder(doc, "Staging Alanındaki Karmaşık Hatalı Veri", "pgAdmin'de 'SELECT * FROM staging.raw_urunler LIMIT 10;' yazın. Sonuçlarda negatif fiyatların, büyük-küçük harf bozukluklarının olduğu o listeyi SS alıp ekleyin.")

    doc.add_heading('4.2. Transform (Temizleme / Dönüştürme)', level=3)
    doc.add_paragraph('İstisnasız her kirli satır tek tek okunur ve mantıksal kontrollerden geçirilir. Tespit edilen kalıtsal yanlışlıklar etl_log isimli log tablomuzdaki "hata_log" grafiğinde sayısal olarak toplanır ki neyi nerede kaybettiğimizi bilelim.')
    
    add_image_placeholder(doc, "ETL hata_log İstatistik Raporu", "pgAdmin'de 'SELECT hata_tipi, COUNT(*) FROM etl_log.hata_log GROUP BY hata_tipi;' çalıştırıp hataların listesini ve miktarını gösteren o sonucun SS'ini alın.")

    doc.add_heading('4.3. Load (Ana Sisteme Yükleme)', level=3)
    doc.add_paragraph('Geçerli olan tüm veriler hedef şemasına taşınır. Burada ON CONFLICT DO UPDATE (UPSERT) teknolojisi kullanılarak veriler ezilmez ve ana veritabanında harika bir yapı oluşur.')

    doc.add_page_break()

    doc.add_heading('5. Sonuç Raporu', level=2)
    doc.add_paragraph('Sürecin tamamı incelendiğinde, tasarlanan bu üç katmanlı (Staging, Transform, Load) ETL mimarisi, dış dünyadan sistemimize sızabilecek her türlü veri kirliliğine karşı aşılmaz bir baraj duvarı işlevi görmektedir. Gerçek dünya senaryolarında, veriler genellikle farklı API uç noktalarından, entegrasyonu tamamlanmamış eski (legacy) veritabanlarından veya veriyi Excel\'e elle giren operasyon çalışanlarının insan kaynaklı (Human Error) anlık dalgınlıklarından dolayı sisteme bozuk formda ulaşır.')
    doc.add_paragraph('Veri biliminde temel bir kural olan "Garbage In, Garbage Out / Çöp Giren, Çöp Çıkar" prensibinden yola çıkarak; eğer bir veritabanı kendi kalitesini denetleyecek bir mekanizmaya sahip değilse, o veri üzerinden üretilecek tüm analizler ve finansal raporlar hatalı sonuçlar verecektir. Kurduğumuz bu otonom veri boru hattı (Data Pipeline), bu hayati riski tamamen sıfıra indirmiştir.')
    doc.add_paragraph('Sistem; negatif fiyatlar, büyük/küçük harf tutarsızlıkları veya geçersiz semboller barındıran formatları insan eli değmesine gerek kalmaksızın tespit edip ana sisteme girişini engellemektedir. Reddedilen hatalı satırlar etl_log.hata_log tablosunda zaman damgasıyla sicile işlenmektedir.')
    doc.add_paragraph('Sonuç olarak; bu projeyle hedef üretim veritabanına ulaşan tüm kayıtlar %100 hijyenik duruma getirilmiştir. Bu yaklaşım firmalara ciddi oranda operasyonel iş gücü ve zaman tasarrufu sağlamaktadır.')
    
    add_image_placeholder(doc, "Temizlenmiş pırıl pırıl Hedef Müşteriler", "pgAdmin'de 'SELECT * FROM hedef.musteriler LIMIT 8;' çalıştırıp tertemiz hale gelmiş veriyi SS alarak kapanış yapın.")

    # Tüm section'lara sayfa Numarası Ekleyelim
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Sayfa X yazılması
        p.add_run("- ")
        add_page_number(p.add_run())
        p.add_run(" -")

    output_path = os.path.join(os.path.dirname(__file__), 'BLM4522_Proje2_ve_5_TEK_RAPOR_v3.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"BİRLEŞTİRİLMİŞ, 10 SAYFALIK, SS YÖNERGELİ TEK RAPOR OLUŞTURULDU: {output_path}")

if __name__ == '__main__':
    generate_unified_report()
